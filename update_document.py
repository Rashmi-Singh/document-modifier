from docx import Document

def parseDocument(document, searchText, replaceText, targetElement):
    performSave = False

    if targetElement == 'para':
        for index, para in enumerate(document.paragraphs):
            if searchText in para.text:
                print('Runs are printed here - \n', para.runs[index].text, '\n')
                performSave = True
                para.runs[index].text = para.runs[index].text.replace(searchText, replaceText)

    if targetElement == 'table':
        for index,table in enumerate(tempdoc1.tables):
            print('\nTable no {}:    {}\n'.format(index+1, table))
            print('Column-wise data -')
            for colIndex, column in enumerate(table.columns):
                for cellIndex, cell in enumerate(column.cells):
                    if (cellIndex > 0):
                        print(cell.text)
                        if searchName in cell.text:
                            cell.text = cell.text.replace(searchName, replaceName)
                            print('after updating text - ', cell.text)
                            performSave = True
                    else:
                        print('Column: {}'.format(cell.text))
                print('\n')

    if performSave:
        print('saving to template1.docx..')
        document.save('./docs/template1.docx')

tempdoc = Document('./docs/template.docx')
searchText = ' for now'
replaceText = '.'
targetElement = 'para'

parseDocument(tempdoc, searchText, replaceText, targetElement) # call for paragraph text change

tempdoc1 = Document('./docs/template1.docx')
searchName = 'Enter file name'
replaceName = 'Sample file'
targetElement = 'table'

parseDocument(tempdoc1, searchName, replaceName, targetElement) # call for table text change
